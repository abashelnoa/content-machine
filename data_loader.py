"""
data_loader.py — טוען את כל הנתונים מהקבצים הקיימים
"""
import os
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).parent

STYLE_GUIDE_PATH = BASE_DIR / "Writing style" / "liraz_style_guide (1).docx"
IDEAS_PATH = BASE_DIR / "רעיונות לפוסטים.docx"
MY_IMAGES_DIR = BASE_DIR / "my images"
STYLE_IMAGES_DIR = BASE_DIR / "style images"

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}

CATEGORY_LABELS = {
    "מוטיבציות": "מוטיבציות",
    "חששות": "חששות",
    "דברים שלא יודעים": "דברים שלא יודעים",
}


def load_style_guide(file_bytes: bytes | None = None, suffix: str = "docx") -> str:
    """קורא את מדריך הסגנון ומחזיר כ-string.
    אם file_bytes ניתן — קורא מזה, אחרת מהקובץ הברירת מחדל."""
    import io
    if file_bytes:
        if suffix == "pdf":
            return _load_style_from_pdf(file_bytes)
        doc = Document(io.BytesIO(file_bytes))
    else:
        doc = Document(STYLE_GUIDE_PATH)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _load_style_from_pdf(file_bytes: bytes) -> str:
    """מחלץ טקסט ממדריך סגנון בפורמט PDF."""
    import io
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(p.strip() for p in pages if p.strip())


def load_post_ideas(file_bytes: bytes | None = None, suffix: str = "docx") -> dict[str, list[str]]:
    """
    קורא את טבלת הרעיונות ומחזיר dict:
    {
        "מוטיבציות": ["רעיון 1", "רעיון 2", ...],
        "חששות": [...],
        "דברים שלא יודעים": [...],
    }
    """
    import io
    if file_bytes:
        if suffix == "xlsx":
            return _load_ideas_from_xlsx(file_bytes)
        if suffix == "pdf":
            return _load_ideas_from_pdf(file_bytes)
        doc = Document(io.BytesIO(file_bytes))
    else:
        doc = Document(IDEAS_PATH)
    ideas: dict[str, list[str]] = {}

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue

            category = cells[0]
            # סנן שורת כותרת
            if category in ("קטגוריה \\ רעיון", "Category \\ Idea", ""):
                continue

            row_ideas = [c for c in cells[1:] if c]
            if row_ideas:
                ideas[category] = row_ideas

    # fallback — אם הטבלה לא נמצאה, קרא מהפסקאות
    if not ideas:
        ideas = _parse_ideas_from_paragraphs(doc)

    return ideas


def _parse_ideas_from_paragraphs(doc: Document) -> dict[str, list[str]]:
    """fallback: קריאת רעיונות מפסקאות אם אין טבלה."""
    result: dict[str, list[str]] = {
        "מוטיבציות": [],
        "חששות": [],
        "דברים שלא יודעים": [],
    }
    current_category = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text in result:
            current_category = text
        elif current_category:
            result[current_category].append(text)
    return result


def _load_ideas_from_pdf(file_bytes: bytes) -> dict[str, list[str]]:
    """מחלץ רעיונות מ-PDF. מצפה לפורמט: שורת כותרת (קטגוריה) ואחריה שורות רעיונות."""
    import io
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    lines = []
    for page in reader.pages:
        text = page.extract_text() or ""
        lines.extend(line.strip() for line in text.splitlines() if line.strip())

    ideas: dict[str, list[str]] = {}
    current_category: str | None = None
    for line in lines:
        # Strip common bullet/numbering prefixes
        clean = line.lstrip("•●▪-–—0123456789.) ").strip()
        if not clean:
            continue
        # Heuristic: a short line with no leading bullet is a category header
        is_bullet = line != clean or len(clean) > 80
        if not is_bullet and len(clean) < 50:
            current_category = clean
            if current_category not in ideas:
                ideas[current_category] = []
        elif current_category and clean:
            ideas[current_category].append(clean)

    return ideas


def _load_ideas_from_xlsx(file_bytes: bytes) -> dict[str, list[str]]:
    """קורא טבלת רעיונות מקובץ XLSX. כל שורה: תא ראשון = קטגוריה, שאר התאים = רעיונות."""
    import io
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    ideas: dict[str, list[str]] = {}
    for row in ws.iter_rows(values_only=True):
        if not row or not row[0]:
            continue
        category = str(row[0]).strip()
        if not category:
            continue
        row_ideas = [str(c).strip() for c in row[1:] if c]
        if row_ideas:
            ideas[category] = row_ideas
    return ideas


def get_my_images() -> list[Path]:
    """מחזיר רשימת נתיבי תמונות אישיות."""
    return sorted(
        p for p in MY_IMAGES_DIR.iterdir()
        if p.suffix.lower() in IMAGE_EXTENSIONS
    )


def get_style_images() -> list[Path]:
    """מחזיר רשימת נתיבי תמונות סגנון."""
    return sorted(
        p for p in STYLE_IMAGES_DIR.iterdir()
        if p.suffix.lower() in IMAGE_EXTENSIONS
    )


def ensure_outputs_dir() -> Path:
    """וודא שתיקיית outputs קיימת."""
    outputs = BASE_DIR / "outputs"
    outputs.mkdir(exist_ok=True)
    return outputs


def create_ideas_docx(ideas_dict: dict) -> bytes:
    """יוצר DOCX מטבלת רעיונות → מחזיר bytes."""
    import io as _io
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()
    title = doc.add_heading("טבלת רעיונות לתוכן", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for audience, ideas in ideas_dict.items():
        doc.add_heading(audience, level=2)
        for i, idea in enumerate(ideas, 1):
            doc.add_paragraph(f"{i}. {idea}")
        doc.add_paragraph()

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_style_guide_docx(content: str) -> bytes:
    """יוצר DOCX ממדריך סגנון כתיבה → מחזיר bytes."""
    import io as _io
    from docx import Document as _Doc

    doc = _Doc()
    doc.add_heading("מדריך סגנון כתיבה", level=1)

    for line in content.split('\n'):
        stripped = line.strip()
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            doc.add_heading(stripped[2:-2], level=2)
        elif stripped:
            doc.add_paragraph(stripped)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
